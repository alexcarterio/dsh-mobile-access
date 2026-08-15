# -*- coding: utf-8 -*-
"""
DSH attachment parser (phone-access companion): convert common formats to
readable text.

Usage:
  py parse_file.py <file> [output]
  (when output is omitted, write to <file>.parsed.txt and print the first
  500 characters to stdout for confirmation)

Supported:
  .txt/.md/.csv/.json/.xml/.yaml/.log and common code files -> read as text
  .docx  -> extract paragraphs with python-docx
  .pdf   -> extract text with pypdf
  .zip   -> extract to a same-named directory with the built-in zipfile
  .7z    -> extract to a same-named directory with py7zr
  .rar   -> extract with rarfile (requires WinRAR/unrar on the system)
  anything else -> error "unsupported format".
"""
import os
import sys
import zipfile

TEXT_EXT = {".txt", ".md", ".csv", ".json", ".xml", ".yml", ".yaml", ".ini",
            ".log", ".js", ".ts", ".py", ".html", ".htm", ".css", ".ps1",
            ".bat", ".sh", ".conf", ".toml", ".sql"}


def read_text(path):
    for enc in ("utf-8", "gbk", "utf-16", "latin-1"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def parse_docx(path):
    import docx
    doc = docx.Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(parts)


def parse_pdf(path):
    from pypdf import PdfReader
    reader = PdfReader(path)
    return "\n\n".join((page.extract_text() or "") for page in reader.pages)


def extract_zip(path, dest):
    with zipfile.ZipFile(path) as z:
        z.extractall(dest)
    return dest


def extract_7z(path, dest):
    import py7zr
    with py7zr.SevenZipFile(path, mode="r") as z:
        z.extractall(dest)
    return dest


def extract_rar(path, dest):
    import rarfile
    with rarfile.RarFile(path) as rf:
        rf.extractall(dest)
    return dest


def main():
    if len(sys.argv) < 2:
        print("Usage: py parse_file.py <file> [output]")
        return 2
    src = sys.argv[1]
    if not os.path.isfile(src):
        print(f"File not found: {src}")
        return 1
    ext = os.path.splitext(src)[1].lower()
    out = sys.argv[2] if len(sys.argv) > 2 else src + ".parsed.txt"
    try:
        if ext in TEXT_EXT:
            text = read_text(src)
        elif ext == ".docx":
            text = parse_docx(src)
        elif ext == ".pdf":
            text = parse_pdf(src)
        elif ext == ".zip":
            dest = os.path.splitext(src)[0]
            os.makedirs(dest, exist_ok=True)
            extract_zip(src, dest)
            with open(out, "w", encoding="utf-8") as f:
                f.write(f"[extracted to] {dest}")
            print(f"[extracted to] {dest}")
            return 0
        elif ext == ".7z":
            dest = os.path.splitext(src)[0]
            os.makedirs(dest, exist_ok=True)
            extract_7z(src, dest)
            with open(out, "w", encoding="utf-8") as f:
                f.write(f"[extracted to] {dest}")
            print(f"[extracted to] {dest}")
            return 0
        elif ext == ".rar":
            dest = os.path.splitext(src)[0]
            os.makedirs(dest, exist_ok=True)
            extract_rar(src, dest)
            with open(out, "w", encoding="utf-8") as f:
                f.write(f"[extracted to] {dest}")
            print(f"[extracted to] {dest}")
            return 0
        else:
            print(f"Unsupported format: {ext} (extractable types: zip/7z/rar)")
            return 1
        with open(out, "w", encoding="utf-8") as f:
            f.write(text)
        print(f"[parsed] {out} ({len(text)} chars)")
        print(text[:500])
        return 0
    except ImportError as e:
        print(f"Missing dependency: {e}. Run: py -3 -m pip install python-docx pypdf py7zr rarfile")
        return 1
    except Exception as e:
        print(f"Parse failed: {e}")
        return 1


if __name__ == "__main__":
    sys.exit(main())
